"""
Professional Document Exporter — v3
Pixel-perfect PDF using ReportLab canvas + platypus for flow
Word (.docx) using python-docx with full style control
"""
import re
from io import BytesIO
from datetime import datetime
from utils.logger import setup_logger

logger = setup_logger(__name__)


# ──────────────────────────────────────────────
# SHARED MARKDOWN PARSER
# ──────────────────────────────────────────────
def parse_lines(content: str):
    tokens = []
    lines  = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith('#### '): tokens.append(('h3', line[5:].strip()))
        elif line.startswith('### '): tokens.append(('h3', line[4:].strip()))
        elif line.startswith('## '):  tokens.append(('h2', line[3:].strip()))
        elif line.startswith('# '):   tokens.append(('h1', line[2:].strip()))

        elif line.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].startswith('|'):
                r = lines[i].strip()
                if not re.match(r'^\|[\s\-:|]+\|$', r):
                    cells = [c.strip() for c in r.split('|')[1:-1]]
                    if any(cells): rows.append(cells)
                i += 1
            if rows: tokens.append(('table', rows))
            continue

        elif re.match(r'^  {0,3}[-*•] ', line):
            tokens.append(('bullet', re.sub(r'^  {0,3}[-*•] ', '', line).strip()))
        elif re.match(r'^ {4,}[-*•] ', line):
            tokens.append(('sub_bullet', re.sub(r'^ +[-*•] ', '', line).strip()))
        elif re.match(r'^\d+\.\s', line):
            tokens.append(('numbered', re.sub(r'^\d+\.\s+', '', line).strip()))
        elif re.match(r'^[-─=]{3,}$', line.strip()):
            tokens.append(('hr', ''))
        elif line.strip() == '':
            tokens.append(('blank', ''))
        else:
            tokens.append(('para', line))

        i += 1
    return tokens


def strip_md(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'`(.+?)`',       r'\1', text)
    return text


def md_to_rl(text):
    """Markdown inline → ReportLab XML-safe string."""
    text = re.sub(r'&(?!(amp|lt|gt|quot|#)\b)', '&amp;', text)
    text = text.replace('<','&lt;').replace('>','&gt;')
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.+?)\*',     r'<i>\1</i>', text)
    text = re.sub(r'`(.+?)`', r'<font name="Courier" color="#c0392b">\1</font>', text)
    return text


# ══════════════════════════════════════════════════════
#  PDF EXPORT  — full canvas + platypus
# ══════════════════════════════════════════════════════
def export_to_pdf(content: str, doc_type: str, department: str, company_name: str = "") -> bytes:
    from reportlab.lib.pagesizes  import A4
    from reportlab.lib.units      import mm, cm
    from reportlab.lib            import colors
    from reportlab.lib.styles     import ParagraphStyle
    from reportlab.lib.enums      import TA_LEFT, TA_CENTER, TA_JUSTIFY, TA_RIGHT
    from reportlab.platypus       import (
        BaseDocTemplate, Frame, PageTemplate,
        Paragraph, Spacer, Table, TableStyle,
        HRFlowable, KeepTogether, NextPageTemplate, PageBreak
    )
    from reportlab.pdfgen         import canvas as rl_canvas

    # ── Dimensions ──
    PW, PH = A4          # 595 × 842 pt
    LM = 20*mm; RM = 20*mm
    TM = 28*mm; BM = 22*mm
    CW = PW - LM - RM    # content width ≈ 170mm

    # ── Palette ──
    NAVY  = colors.HexColor('#1B3A6B')
    BLUE  = colors.HexColor('#2563AE')
    TEAL  = colors.HexColor('#0D7A6E')
    DARK  = colors.HexColor('#1A1A2E')
    BODY  = colors.HexColor('#2D2D2D')
    GREY  = colors.HexColor('#5A5A72')
    LGREY = colors.HexColor('#9A9AB0')
    SMOKE = colors.HexColor('#F4F6FB')
    STRIPE= colors.HexColor('#EAF1FB')
    WHITE = colors.white

    FONT  = 'Helvetica'
    FONTB = 'Helvetica-Bold'
    FONTI = 'Helvetica-Oblique'

    # ── Styles ──
    def S(name, **kw):
        return ParagraphStyle(name, **kw)

    sH1 = S('H1', fontName=FONTB, fontSize=15, textColor=NAVY,
             leading=20, spaceBefore=16, spaceAfter=4, alignment=TA_LEFT)
    sH2 = S('H2', fontName=FONTB, fontSize=12.5, textColor=BLUE,
             leading=17, spaceBefore=12, spaceAfter=3, alignment=TA_LEFT)
    sH3 = S('H3', fontName=FONTB, fontSize=11, textColor=TEAL,
             leading=15, spaceBefore=8,  spaceAfter=2, alignment=TA_LEFT)
    sBd = S('Bd', fontName=FONT,  fontSize=10, textColor=BODY,
             leading=15, spaceBefore=0,  spaceAfter=4, alignment=TA_JUSTIFY)
    sBu = S('Bu', fontName=FONT,  fontSize=10, textColor=BODY,
             leading=14, spaceBefore=1,  spaceAfter=1,
             leftIndent=12, bulletIndent=0)
    sSb = S('Sb', fontName=FONT,  fontSize=9.5, textColor=GREY,
             leading=13, spaceBefore=0,  spaceAfter=1,
             leftIndent=24, bulletIndent=0)
    sNm = S('Nm', fontName=FONT,  fontSize=10, textColor=BODY,
             leading=14, spaceBefore=1,  spaceAfter=1,
             leftIndent=12)
    sTH = S('TH', fontName=FONTB, fontSize=9,  textColor=WHITE,
             leading=12, alignment=TA_LEFT)
    sTC = S('TC', fontName=FONT,  fontSize=9,  textColor=DARK,
             leading=12, alignment=TA_LEFT)

    # ── Page template with header/footer ──
    class DocCanvas(BaseDocTemplate):
        def __init__(self, buf, **kw):
            super().__init__(buf, **kw)
            frame = Frame(LM, BM, CW, PH - TM - BM, id='main',
                          leftPadding=0, rightPadding=0,
                          topPadding=0,  bottomPadding=0)
            self.addPageTemplates([
                PageTemplate(id='cover', frames=[frame],
                             onPage=self._cover_page),
                PageTemplate(id='body',  frames=[frame],
                             onPage=self._body_page),
            ])

        def _draw_header(self, c):
            c.saveState()
            # Top color bar
            c.setFillColor(NAVY)
            c.rect(0, PH - 14*mm, PW, 14*mm, fill=True, stroke=False)
            # Title in bar
            c.setFillColor(WHITE)
            c.setFont(FONTB, 9)
            c.drawString(LM, PH - 9*mm, f'{doc_type}  │  {department}')
            c.setFont(FONT, 8)
            c.setFillColor(colors.HexColor('#A8C4E8'))
            c.drawRightString(PW - RM, PH - 9*mm, company_name or '')
            c.restoreState()

        def _draw_footer(self, c, page_num):
            c.saveState()
            y = BM - 8*mm
            # Footer line
            c.setStrokeColor(colors.HexColor('#C8D4E8'))
            c.setLineWidth(0.5)
            c.line(LM, y + 5*mm, PW - RM, y + 5*mm)
            # Footer text
            c.setFont(FONT, 7.5)
            c.setFillColor(LGREY)
            date_str = datetime.now().strftime('%d %B %Y')
            c.drawString(LM, y + 1.5*mm, f'{date_str}  ·  Confidential')
            c.drawRightString(PW - RM, y + 1.5*mm, f'Page {page_num}')
            c.restoreState()

        def _cover_page(self, c, doc):
            self._draw_footer(c, doc.page)

        def _body_page(self, c, doc):
            self._draw_header(c)
            self._draw_footer(c, doc.page)

    buf = BytesIO()
    pdf = DocCanvas(
        buf,
        pagesize=A4,
        leftMargin=LM, rightMargin=RM,
        topMargin=TM,  bottomMargin=BM,
        title=f'{doc_type} — {department}',
        author=company_name or 'DocForgeHub',
    )

    story = []

    # ══════════════════
    # COVER PAGE
    # ══════════════════
    story.append(NextPageTemplate('cover'))

    # Big colored cover block drawn via canvas hack using Table
    cover_data = [[Paragraph(
        f'<font color="#FFFFFF"><b><font size="28">{doc_type.upper()}</font></b><br/>'
        f'<font size="14" color="#A8C4E8">{department}</font><br/><br/>'
        f'<font size="10" color="#7BA8D4">'
        f'{company_name + "   ·   " if company_name else ""}'
        f'{datetime.now().strftime("%B %d, %Y")}</font></font>',
        S('Cov', fontName=FONTB, fontSize=28, textColor=WHITE,
          leading=36, alignment=TA_LEFT))
    ]]
    cover_tbl = Table(cover_data, colWidths=[CW])
    cover_tbl.setStyle(TableStyle([
        ('BACKGROUND',    (0,0), (-1,-1), NAVY),
        ('LEFTPADDING',   (0,0), (-1,-1), 18),
        ('RIGHTPADDING',  (0,0), (-1,-1), 18),
        ('TOPPADDING',    (0,0), (-1,-1), 22),
        ('BOTTOMPADDING', (0,0), (-1,-1), 22),
        ('ROUNDEDCORNERS',(0,0), (-1,-1), [4,4,4,4]),
    ]))
    story.append(Spacer(1, 8*mm))
    story.append(cover_tbl)
    story.append(Spacer(1, 6*mm))

    # Accent rule
    story.append(HRFlowable(width=CW, thickness=2, color=BLUE,
                             spaceBefore=2, spaceAfter=10))

    # "About this document" summary box
    about_data = [[
        Paragraph(f'<b>Document Type:</b> {doc_type}', sTC),
        Paragraph(f'<b>Department:</b> {department}', sTC),
        Paragraph(f'<b>Date:</b> {datetime.now().strftime("%Y-%m-%d")}', sTC),
    ]]
    about_tbl = Table(about_data, colWidths=[CW/3]*3)
    about_tbl.setStyle(TableStyle([
        ('BACKGROUND',    (0,0), (-1,-1), SMOKE),
        ('LEFTPADDING',   (0,0), (-1,-1), 10),
        ('RIGHTPADDING',  (0,0), (-1,-1), 10),
        ('TOPPADDING',    (0,0), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
        ('LINEABOVE',     (0,0), (-1,0),  1, BLUE),
        ('LINEBELOW',     (0,0), (-1,-1), 1, BLUE),
        ('VALIGN',        (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story.append(about_tbl)
    story.append(Spacer(1, 10*mm))
    story.append(HRFlowable(width=CW, thickness=0.5,
                             color=colors.HexColor('#D0DCF0'),
                             spaceBefore=2, spaceAfter=8))

    # Switch to body pages with header
    story.append(NextPageTemplate('body'))
    story.append(PageBreak())

    # ══════════════════
    # CONTENT
    # ══════════════════
    tokens  = parse_lines(content)
    num_ctr = 0

    i = 0
    while i < len(tokens):
        typ, dat = tokens[i]

        # ── H1 ──
        if typ == 'h1':
            num_ctr = 0
            block = [
                HRFlowable(width=CW, thickness=1, color=NAVY,
                           spaceBefore=10, spaceAfter=3),
                Paragraph(md_to_rl(dat), sH1),
                HRFlowable(width=CW, thickness=0.4,
                           color=colors.HexColor('#C0D0EA'),
                           spaceBefore=2, spaceAfter=8),
            ]
            story.append(KeepTogether(block))

        # ── H2 ──
        elif typ == 'h2':
            num_ctr = 0
            story.append(KeepTogether([
                Spacer(1, 4),
                Paragraph(md_to_rl(dat), sH2),
                HRFlowable(width=CW * 0.55, thickness=0.5,
                           color=colors.HexColor('#BAD0EE'),
                           spaceBefore=1, spaceAfter=5),
            ]))

        # ── H3 ──
        elif typ == 'h3':
            story.append(Paragraph(md_to_rl(dat), sH3))

        # ── BULLET ──
        elif typ == 'bullet':
            num_ctr = 0
            story.append(Paragraph(f'<font color="#2563AE">▸</font>&nbsp;&nbsp;{md_to_rl(dat)}', sBu))

        # ── SUB-BULLET ──
        elif typ == 'sub_bullet':
            story.append(Paragraph(f'<font color="#9A9AB0">◦</font>&nbsp;&nbsp;{md_to_rl(dat)}', sSb))

        # ── NUMBERED ──
        elif typ == 'numbered':
            num_ctr += 1
            story.append(Paragraph(
                f'<font color="#1B3A6B"><b>{num_ctr}.</b></font>&nbsp;&nbsp;{md_to_rl(dat)}',
                sNm))

        # ── HR ──
        elif typ == 'hr':
            story.append(HRFlowable(width=CW, thickness=0.5,
                                    color=colors.HexColor('#D8DFF0'),
                                    spaceBefore=5, spaceAfter=5))

        # ── BLANK ──
        elif typ == 'blank':
            story.append(Spacer(1, 3))

        # ── TABLE ──
        elif typ == 'table':
            rows     = dat
            num_cols = max(len(r) for r in rows)
            if not rows or num_cols == 0:
                i += 1; continue

            # Normalise column count
            norm = []
            for r in rows:
                row = list(r) + [''] * (num_cols - len(r))
                norm.append(row[:num_cols])

            # Smart column widths: first col slightly wider if needed
            if num_cols == 1:
                col_ws = [CW]
            else:
                base = CW / num_cols
                col_ws = [base] * num_cols

            # Build cell paragraphs
            tbl_data = []
            for ri, row in enumerate(norm):
                tbl_data.append([
                    Paragraph(md_to_rl(cell),
                               sTH if ri == 0 else sTC)
                    for cell in row
                ])

            tbl = Table(tbl_data, colWidths=col_ws, repeatRows=1, hAlign='LEFT')

            ts = [
                # Header row
                ('BACKGROUND',    (0,0),  (-1,0),  NAVY),
                ('TEXTCOLOR',     (0,0),  (-1,0),  WHITE),
                ('FONTNAME',      (0,0),  (-1,0),  FONTB),
                ('FONTSIZE',      (0,0),  (-1,-1), 9),
                ('LINEBELOW',     (0,0),  (-1,0),  1.5, BLUE),

                # All cells
                ('FONTNAME',      (0,1),  (-1,-1), FONT),
                ('VALIGN',        (0,0),  (-1,-1), 'MIDDLE'),
                ('LEFTPADDING',   (0,0),  (-1,-1), 7),
                ('RIGHTPADDING',  (0,0),  (-1,-1), 7),
                ('TOPPADDING',    (0,0),  (-1,-1), 5),
                ('BOTTOMPADDING', (0,0),  (-1,-1), 5),

                # Borders
                ('INNERGRID',     (0,0),  (-1,-1), 0.3, colors.HexColor('#C4D4E8')),
                ('BOX',           (0,0),  (-1,-1), 0.5, colors.HexColor('#A0BCE0')),
            ]

            # Alternating row colours
            for ri in range(1, len(norm)):
                bg = STRIPE if ri % 2 == 1 else WHITE
                ts.append(('BACKGROUND', (0, ri), (-1, ri), bg))

            tbl.setStyle(TableStyle(ts))
            story.append(Spacer(1, 5))
            story.append(KeepTogether([tbl, Spacer(1, 8)]))

        # ── PARAGRAPH ──
        elif typ == 'para':
            if dat.strip():
                num_ctr = 0
                story.append(Paragraph(md_to_rl(dat), sBd))

        i += 1

    pdf.build(story)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════
#  WORD (.docx) EXPORT
# ══════════════════════════════════════════════════════
def export_to_docx(content: str, doc_type: str, department: str, company_name: str = "") -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Twips, Cm
    from docx.enum.text  import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns    import qn
    from docx.oxml       import OxmlElement

    # Colours
    NAVY   = RGBColor(0x1B, 0x3A, 0x6B)
    BLUE   = RGBColor(0x25, 0x63, 0xAE)
    TEAL   = RGBColor(0x0D, 0x7A, 0x6E)
    DARK   = RGBColor(0x1A, 0x1A, 0x2E)
    BODY   = RGBColor(0x2D, 0x2D, 0x2D)
    GREY   = RGBColor(0x5A, 0x5A, 0x72)
    LGREY  = RGBColor(0x9A, 0x9A, 0xB0)
    WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
    NAVY_X = '1B3A6B'
    BLUE_X = '2563AE'
    SMOK_X = 'F4F6FB'
    STRI_X = 'EAF1FB'
    WHIT_X = 'FFFFFF'

    FONT = 'Calibri'
    doc  = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_width    = Inches(8.27)   # A4
    sec.page_height   = Inches(11.69)
    sec.left_margin   = Inches(1.0)
    sec.right_margin  = Inches(1.0)
    sec.top_margin    = Inches(1.0)
    sec.bottom_margin = Inches(0.9)
    CW_IN = 8.27 - 1.0 - 1.0   # 6.27"
    CW_TW = int(CW_IN * 1440)   # twips

    # ── Helpers ──
    def cell_shading(cell, hex_fill):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for x in tcPr.findall(qn('w:shd')): tcPr.remove(x)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_fill)
        tcPr.append(shd)

    def bottom_border(para, color_hex, size=6):
        pPr  = para._p.get_or_add_pPr()
        for x in pPr.findall(qn('w:pBdr')): pPr.remove(x)
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    str(size))
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), color_hex)
        pBdr.append(bot)
        pPr.append(pBdr)

    def para_spacing(para, before=0, after=6, line=None):
        pf = para.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after  = Pt(after)
        if line:
            from docx.enum.text import WD_LINE_SPACING
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(line)

    def add_runs(para, text, size=11, color=None, bold=False):
        if color is None: color = BODY
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
        for p in parts:
            if p.startswith('**') and p.endswith('**'):
                r = para.add_run(p[2:-2])
                r.bold = True
                r.font.color.rgb = color
            elif p.startswith('*') and p.endswith('*'):
                r = para.add_run(p[1:-1])
                r.italic = True
                r.font.color.rgb = color
            elif p.startswith('`') and p.endswith('`'):
                r = para.add_run(p[1:-1])
                r.font.name = 'Courier New'
                r.font.size = Pt(size - 0.5)
                r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            elif p:
                r = para.add_run(p)
                r.font.color.rgb = color
            if p:
                r.font.name = FONT if not (p.startswith('`') and p.endswith('`')) else 'Courier New'
                r.font.size = Pt(size)
                if bold and not (p.startswith('*') or p.startswith('`')):
                    r.bold = True

    # ── Cover Page ──
    # Top heavy rule
    cp = doc.add_paragraph()
    para_spacing(cp, 0, 20)
    tc_elem = cp._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top_el = OxmlElement('w:top')
    top_el.set(qn('w:val'),   'single')
    top_el.set(qn('w:sz'),    '36')
    top_el.set(qn('w:space'), '1')
    top_el.set(qn('w:color'), NAVY_X)
    pBdr.append(top_el)
    tc_elem.append(pBdr)

    # Document type large
    pt = doc.add_paragraph()
    para_spacing(pt, 4, 2)
    r = pt.add_run(doc_type.upper())
    r.font.name  = FONT
    r.font.size  = Pt(30)
    r.font.bold  = True
    r.font.color.rgb = NAVY

    # Department
    pd = doc.add_paragraph()
    para_spacing(pd, 0, 2)
    r = pd.add_run(department)
    r.font.name  = FONT
    r.font.size  = Pt(15)
    r.font.color.rgb = BLUE

    # Meta line
    pm = doc.add_paragraph()
    para_spacing(pm, 0, 16)
    meta_str = (f'{company_name}  ·  ' if company_name else '') + datetime.now().strftime('%B %d, %Y')
    r = pm.add_run(meta_str)
    r.font.name  = FONT
    r.font.size  = Pt(10)
    r.font.color.rgb = LGREY

    # Thick rule
    pdiv = doc.add_paragraph()
    bottom_border(pdiv, NAVY_X, 18)
    para_spacing(pdiv, 0, 18)

    # Info table (cover summary)
    info_tbl = doc.add_table(rows=1, cols=3)
    info_tbl.style = 'Table Grid'
    info_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    info_w = CW_TW // 3
    labels = [
        ('Document Type', doc_type),
        ('Department', department),
        ('Date', datetime.now().strftime('%Y-%m-%d')),
    ]
    for ci, (lbl, val) in enumerate(labels):
        cell = info_tbl.rows[0].cells[ci]
        cell_shading(cell, SMOK_X)
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW  = OxmlElement('w:tcW')
        tcW.set(qn('w:w'),    str(info_w))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        p = cell.paragraphs[0]
        para_spacing(p, 4, 4)
        r1 = p.add_run(lbl + ': ')
        r1.bold = True; r1.font.name = FONT; r1.font.size = Pt(9); r1.font.color.rgb = NAVY
        r2 = p.add_run(val)
        r2.font.name = FONT; r2.font.size = Pt(9); r2.font.color.rgb = BODY

    doc.add_paragraph()

    # ── Content ──
    tokens  = parse_lines(content)
    num_ctr = 0

    for typ, dat in tokens:

        if typ == 'h1':
            num_ctr = 0
            p = doc.add_paragraph()
            para_spacing(p, 18, 4)
            add_runs(p, dat, size=17, color=NAVY, bold=True)
            bottom_border(p, BLUE_X, 8)

        elif typ == 'h2':
            num_ctr = 0
            p = doc.add_paragraph()
            para_spacing(p, 13, 3)
            add_runs(p, dat, size=13, color=BLUE, bold=True)
            bottom_border(p, 'C0D4EE', 4)

        elif typ == 'h3':
            p = doc.add_paragraph()
            para_spacing(p, 9, 2)
            add_runs(p, dat, size=11, color=TEAL, bold=True)

        elif typ == 'bullet':
            num_ctr = 0
            p = doc.add_paragraph(style='List Bullet')
            para_spacing(p, 1, 2)
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, dat, size=10.5, color=BODY)

        elif typ == 'sub_bullet':
            p = doc.add_paragraph(style='List Bullet 2')
            para_spacing(p, 0, 1)
            p.paragraph_format.left_indent = Inches(0.55)
            add_runs(p, dat, size=10, color=GREY)

        elif typ == 'numbered':
            num_ctr += 1
            p = doc.add_paragraph(style='List Number')
            para_spacing(p, 1, 2)
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, dat, size=10.5, color=BODY)

        elif typ == 'hr':
            p = doc.add_paragraph()
            bottom_border(p, 'D0DCF0', 4)
            para_spacing(p, 4, 4)

        elif typ == 'blank':
            p = doc.add_paragraph()
            para_spacing(p, 0, 3)

        elif typ == 'table':
            rows     = dat
            num_cols = max(len(r) for r in rows)
            if not rows or num_cols == 0: continue

            col_w_tw = CW_TW // num_cols
            norm = [(list(r) + [''] * (num_cols - len(r)))[:num_cols] for r in rows]

            tbl = doc.add_table(rows=len(norm), cols=num_cols)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Set table width
            tPr_found = tbl._tbl.find(qn('w:tblPr'))
            tPr  = tPr_found if tPr_found is not None else OxmlElement('w:tblPr')
            tW   = OxmlElement('w:tblW')
            tW.set(qn('w:w'), str(CW_TW))
            tW.set(qn('w:type'), 'dxa')
            tPr.append(tW)

            for ri, row_data in enumerate(norm):
                is_hdr = (ri == 0)
                tr = tbl.rows[ri]
                for ci, cell_text in enumerate(row_data):
                    cell = tr.cells[ci]
                    text = strip_md(cell_text)

                    # Width
                    tc   = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW  = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'),    str(col_w_tw))
                    tcW.set(qn('w:type'), 'dxa')
                    tcPr.append(tcW)

                    # Vertical align
                    vAl = OxmlElement('w:vAlign')
                    vAl.set(qn('w:val'), 'center')
                    tcPr.append(vAl)

                    # Cell padding
                    tcMar = OxmlElement('w:tcMar')
                    for side, tw in [('top','80'),('bottom','80'),('left','110'),('right','110')]:
                        m = OxmlElement(f'w:{side}')
                        m.set(qn('w:w'),    tw)
                        m.set(qn('w:type'), 'dxa')
                        tcMar.append(m)
                    tcPr.append(tcMar)

                    # Background
                    if is_hdr:
                        cell_shading(cell, NAVY_X)
                    elif ri % 2 == 1:
                        cell_shading(cell, STRI_X)
                    else:
                        cell_shading(cell, WHIT_X)

                    # Text
                    p = cell.paragraphs[0]
                    para_spacing(p, 0, 0)
                    r = p.add_run(text)
                    r.font.name  = FONT
                    r.font.size  = Pt(9.5)
                    r.font.bold  = is_hdr
                    r.font.color.rgb = WHITE if is_hdr else BODY

            doc.add_paragraph()

        elif typ == 'para':
            if not dat.strip(): continue
            num_ctr = 0
            p = doc.add_paragraph()
            para_spacing(p, 0, 5)
            p.paragraph_format.line_spacing = Pt(14.5)
            add_runs(p, dat, size=10.5, color=BODY)

    # ── Header ──
    hdr = sec.header
    hdr.is_linked_to_previous = False
    hp  = hdr.paragraphs[0]
    hp.clear(); hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(f'{doc_type}  │  {department}')
    r.font.name = FONT; r.font.size = Pt(8); r.font.color.rgb = LGREY
    bottom_border(hp, 'C8D4E8', 3)

    # ── Footer ──
    ftr = sec.footer
    ftr.is_linked_to_previous = False
    fp  = ftr.paragraphs[0]
    fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    top_bdr = fp._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    t2 = OxmlElement('w:top')
    t2.set(qn('w:val'), 'single'); t2.set(qn('w:sz'), '3')
    t2.set(qn('w:space'), '1');    t2.set(qn('w:color'), 'C8D4E8')
    pBdr2.append(t2); top_bdr.append(pBdr2)

    date_str = datetime.now().strftime('%Y-%m-%d')
    left_r = fp.add_run(f'{date_str}  ·  Confidential    Page ')
    left_r.font.name = FONT; left_r.font.size = Pt(8); left_r.font.color.rgb = LGREY

    # Page number field
    for tag, txt in [('begin',None),('separate',None),('end',None)]:
        fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), tag)
        if tag == 'begin':
            instr = OxmlElement('w:instrText')
            instr.text = ' PAGE '
            instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        rr = fp.add_run()
        rr.font.name = FONT; rr.font.size = Pt(8); rr.font.color.rgb = LGREY
        rr._r.append(fc)
        if tag == 'begin': rr._r.append(instr)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
